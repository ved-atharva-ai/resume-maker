import os
os.environ['GRPC_VERBOSITY'] = 'ERROR'
os.environ['GLOG_minloglevel'] = '2'

import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import PyPDF2
import time

# Page configuration
st.set_page_config(page_title="AI Resume Optimizer", layout="wide", page_icon="📄")

# Initialize session state
if 'optimized_resume' not in st.session_state:
    st.session_state.optimized_resume = None

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX"""
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return None

# Resume extraction prompt
extraction_prompt = PromptTemplate(
    input_variables=["resume_text"],
    template="""Extract all information from this resume and structure it in JSON format.

Resume Text:
{resume_text}

Extract the following information and return ONLY a valid JSON object:
{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "summary": "Professional summary or objective",
    "skills": ["skill1", "skill2", ...],
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "duration": "Start Date - End Date",
            "responsibilities": ["responsibility1", "responsibility2", ...]
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "university": "University Name",
            "year": "Graduation Year",
            "details": "Additional details if any"
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Project Description",
            "technologies": ["tech1", "tech2", ...],
            "duration": "Duration or Date"
        }}
    ],
    "certifications": ["cert1", "cert2", ...]
}}

If any section is not found, use an empty array [] or empty string "". Return ONLY valid JSON, no additional text."""
)

# Resume optimization prompt
optimization_prompt = PromptTemplate(
    input_variables=["extracted_resume", "job_requirements"],
    template="""You are an expert resume writer. Optimize this resume based on the job requirements.

Current Resume Data:
{extracted_resume}

Job Requirements:
{job_requirements}

CRITICAL RULES:
1. Keep Education section EXACTLY as is - DO NOT modify
2. Keep Projects section EXACTLY as is - DO NOT modify
3. Keep Name, Email, Phone EXACTLY as is - DO NOT modify
4. You CAN modify:
   - Professional Summary (tailor to job requirements)
   - Skills (reorder and highlight relevant skills, add missing relevant skills)
   - Work Experience (rewrite descriptions to match job requirements, use action verbs and metrics)
   - Certifications (reorder by relevance, suggest relevant ones if needed)

Generate an optimized resume that:
- Highlights relevant experience matching the job requirements
- Uses keywords from the job description
- Quantifies achievements where possible
- Uses strong action verbs
- Maintains professional tone
- Keeps all education and project information unchanged

Return ONLY a valid JSON object with the same structure as the input, with optimized content."""
)

def generate_stylish_pdf(resume_data):
    """Generate a modern, stylish PDF resume"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter, 
        topMargin=0.5*inch, 
        bottomMargin=0.5*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    story = []
    
    # Define modern color scheme
    primary_color = HexColor('#1a5490')  # Professional blue
    secondary_color = HexColor('#2c3e50')  # Dark blue-gray
    accent_color = HexColor('#3498db')  # Light blue
    text_color = HexColor('#2c3e50')  # Dark text
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Name style - Large and bold
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=primary_color,
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        leading=32
    )
    
    # Contact style
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        textColor=secondary_color,
        alignment=TA_CENTER,
        spaceAfter=16,
        leading=14
    )
    
    # Section heading style
    section_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=primary_color,
        spaceAfter=10,
        spaceBefore=14,
        fontName='Helvetica-Bold',
        borderWidth=0,
        borderColor=primary_color,
        borderPadding=0,
        leftIndent=0,
        leading=16
    )
    
    # Body text style
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        textColor=text_color,
        spaceAfter=8,
        leading=14,
        alignment=TA_JUSTIFY
    )
    
    # Job title style
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=11,
        textColor=secondary_color,
        fontName='Helvetica-Bold',
        spaceAfter=2,
        leading=14
    )
    
    # Company style
    company_style = ParagraphStyle(
        'Company',
        parent=styles['Normal'],
        fontSize=10,
        textColor=accent_color,
        fontName='Helvetica-Bold',
        spaceAfter=2,
        leading=12
    )
    
    # Duration style
    duration_style = ParagraphStyle(
        'Duration',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#7f8c8d'),
        fontName='Helvetica-Oblique',
        spaceAfter=6,
        leading=11
    )
    
    # Bullet style
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        textColor=text_color,
        leftIndent=20,
        spaceAfter=4,
        leading=13,
        bulletIndent=10
    )
    
    # Skills style
    skills_style = ParagraphStyle(
        'Skills',
        parent=styles['Normal'],
        fontSize=10,
        textColor=text_color,
        spaceAfter=6,
        leading=14
    )
    
    # === HEADER ===
    # Name
    story.append(Paragraph(resume_data['name'].upper(), name_style))
    
    # Contact Info
    contact_parts = []
    if resume_data.get('email'):
        contact_parts.append(resume_data['email'])
    if resume_data.get('phone'):
        contact_parts.append(resume_data['phone'])
    
    if contact_parts:
        contact_text = " | ".join(contact_parts)
        story.append(Paragraph(contact_text, contact_style))
    
    # Decorative line
    story.append(HRFlowable(
        width="100%", 
        thickness=2, 
        color=primary_color, 
        spaceAfter=16,
        spaceBefore=0
    ))
    
    # === PROFESSIONAL SUMMARY ===
    if resume_data.get('summary'):
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading_style))
        story.append(Paragraph(resume_data['summary'], body_style))
        story.append(Spacer(1, 0.15*inch))
    
    # === SKILLS ===
    if resume_data.get('skills'):
        story.append(Paragraph("CORE COMPETENCIES", section_heading_style))
        
        # Format skills in a clean way
        skills_text = " • ".join(resume_data['skills'])
        story.append(Paragraph(skills_text, skills_style))
        story.append(Spacer(1, 0.15*inch))
    
    # === WORK EXPERIENCE ===
    if resume_data.get('experience'):
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_heading_style))
        
        for exp in resume_data['experience']:
            # Job title
            story.append(Paragraph(exp.get('title', 'Position'), job_title_style))
            
            # Company name
            story.append(Paragraph(exp.get('company', 'Company'), company_style))
            
            # Duration
            if exp.get('duration'):
                story.append(Paragraph(exp['duration'], duration_style))
            
            # Responsibilities
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    bullet_text = f"• {resp}"
                    story.append(Paragraph(bullet_text, bullet_style))
            
            story.append(Spacer(1, 0.12*inch))
    
    # === PROJECTS ===
    if resume_data.get('projects'):
        story.append(Paragraph("PROJECTS", section_heading_style))
        
        for project in resume_data['projects']:
            # Project name
            project_name = project.get('name', 'Project')
            if project.get('duration'):
                project_name += f" ({project['duration']})"
            story.append(Paragraph(project_name, job_title_style))
            
            # Description
            if project.get('description'):
                story.append(Paragraph(project['description'], body_style))
            
            # Technologies
            if project.get('technologies'):
                tech_text = f"<b>Technologies:</b> {', '.join(project['technologies'])}"
                story.append(Paragraph(tech_text, skills_style))
            
            story.append(Spacer(1, 0.1*inch))
    
    # === EDUCATION ===
    if resume_data.get('education'):
        story.append(Paragraph("EDUCATION", section_heading_style))
        
        education_list = resume_data['education'] if isinstance(resume_data['education'], list) else [resume_data['education']]
        
        for edu in education_list:
            # Degree
            degree_text = f"<b>{edu.get('degree', 'Degree')}</b>"
            story.append(Paragraph(degree_text, job_title_style))
            
            # University
            story.append(Paragraph(edu.get('university', 'University'), company_style))
            
            # Year and details
            year_details = []
            if edu.get('year'):
                year_details.append(str(edu['year']))
            if edu.get('details'):
                year_details.append(edu['details'])
            
            if year_details:
                story.append(Paragraph(" | ".join(year_details), duration_style))
            
            story.append(Spacer(1, 0.08*inch))
    
    # === CERTIFICATIONS ===
    if resume_data.get('certifications'):
        story.append(Paragraph("CERTIFICATIONS", section_heading_style))
        
        for cert in resume_data['certifications']:
            cert_text = f"• {cert}"
            story.append(Paragraph(cert_text, bullet_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_stylish_docx(resume_data):
    """Generate a modern, stylish DOCX resume"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Define colors
    primary_color = RGBColor(26, 84, 144)  # Professional blue
    secondary_color = RGBColor(44, 62, 80)  # Dark blue-gray
    accent_color = RGBColor(52, 152, 219)  # Light blue
    
    # === NAME ===
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(resume_data['name'].upper())
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = primary_color
    
    # === CONTACT INFO ===
    contact_parts = []
    if resume_data.get('email'):
        contact_parts.append(resume_data['email'])
    if resume_data.get('phone'):
        contact_parts.append(resume_data['phone'])
    
    if contact_parts:
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_paragraph.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = secondary_color
    
    # Add horizontal line
    doc.add_paragraph("_" * 80)
    
    # === PROFESSIONAL SUMMARY ===
    if resume_data.get('summary'):
        add_section_heading(doc, "PROFESSIONAL SUMMARY", primary_color)
        summary_para = doc.add_paragraph(resume_data['summary'])
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        format_body_text(summary_para)
        doc.add_paragraph()
    
    # === SKILLS ===
    if resume_data.get('skills'):
        add_section_heading(doc, "CORE COMPETENCIES", primary_color)
        skills_text = " • ".join(resume_data['skills'])
        skills_para = doc.add_paragraph(skills_text)
        format_body_text(skills_para)
        doc.add_paragraph()
    
    # === WORK EXPERIENCE ===
    if resume_data.get('experience'):
        add_section_heading(doc, "PROFESSIONAL EXPERIENCE", primary_color)
        
        for exp in resume_data['experience']:
            # Job title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', 'Position'))
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            title_run.font.color.rgb = secondary_color
            
            # Company
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(exp.get('company', 'Company'))
            company_run.font.size = Pt(10)
            company_run.font.bold = True
            company_run.font.color.rgb = accent_color
            
            # Duration
            if exp.get('duration'):
                duration_para = doc.add_paragraph()
                duration_run = duration_para.add_run(exp['duration'])
                duration_run.font.size = Pt(9)
                duration_run.font.italic = True
                duration_run.font.color.rgb = RGBColor(127, 140, 141)
            
            # Responsibilities
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    resp_para = doc.add_paragraph(resp, style='List Bullet')
                    format_body_text(resp_para)
            
            doc.add_paragraph()
    
    # === PROJECTS ===
    if resume_data.get('projects'):
        add_section_heading(doc, "PROJECTS", primary_color)
        
        for project in resume_data['projects']:
            # Project name
            project_name = project.get('name', 'Project')
            if project.get('duration'):
                project_name += f" ({project['duration']})"
            
            project_para = doc.add_paragraph()
            project_run = project_para.add_run(project_name)
            project_run.font.size = Pt(11)
            project_run.font.bold = True
            project_run.font.color.rgb = secondary_color
            
            # Description
            if project.get('description'):
                desc_para = doc.add_paragraph(project['description'])
                desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                format_body_text(desc_para)
            
            # Technologies
            if project.get('technologies'):
                tech_para = doc.add_paragraph()
                tech_label = tech_para.add_run("Technologies: ")
                tech_label.font.bold = True
                tech_label.font.size = Pt(10)
                tech_text = tech_para.add_run(", ".join(project['technologies']))
                tech_text.font.size = Pt(10)
            
            doc.add_paragraph()
    
    # === EDUCATION ===
    if resume_data.get('education'):
        add_section_heading(doc, "EDUCATION", primary_color)
        
        education_list = resume_data['education'] if isinstance(resume_data['education'], list) else [resume_data['education']]
        
        for edu in education_list:
            # Degree
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get('degree', 'Degree'))
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            degree_run.font.color.rgb = secondary_color
            
            # University
            uni_para = doc.add_paragraph()
            uni_run = uni_para.add_run(edu.get('university', 'University'))
            uni_run.font.size = Pt(10)
            uni_run.font.bold = True
            uni_run.font.color.rgb = accent_color
            
            # Year and details
            year_details = []
            if edu.get('year'):
                year_details.append(str(edu['year']))
            if edu.get('details'):
                year_details.append(edu['details'])
            
            if year_details:
                year_para = doc.add_paragraph()
                year_run = year_para.add_run(" | ".join(year_details))
                year_run.font.size = Pt(9)
                year_run.font.italic = True
                year_run.font.color.rgb = RGBColor(127, 140, 141)
            
            doc.add_paragraph()
    
    # === CERTIFICATIONS ===
    if resume_data.get('certifications'):
        add_section_heading(doc, "CERTIFICATIONS", primary_color)
        
        for cert in resume_data['certifications']:
            cert_para = doc.add_paragraph(cert, style='List Bullet')
            format_body_text(cert_para)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def add_section_heading(doc, text, color):
    """Add a styled section heading to the document"""
    heading = doc.add_paragraph()
    heading_run = heading.add_run(text)
    heading_run.font.size = Pt(13)
    heading_run.font.bold = True
    heading_run.font.color.rgb = color

def format_body_text(paragraph):
    """Format body text paragraphs"""
    for run in paragraph.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(44, 62, 80)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .step-container {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1rem;
        border-left: 4px solid #667eea;
    }
    .stButton>button {
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem;
        font-weight: 600;
        border-radius: 8px;
    }
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .download-section {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin: 2rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Main UI
st.markdown('<div class="main-header"><h1>🎯 AI-Powered Resume Optimizer</h1><p>Upload your resume and job requirements to get a tailored, professional resume</p></div>', unsafe_allow_html=True)

# Sidebar for API Key
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Google Gemini API Key", type="password", help="Enter your Gemini API key")
    
    st.markdown("---")
    st.markdown("### 📋 How it works:")
    st.markdown("""
    1. Upload your current resume (PDF/DOCX)
    2. Paste the job requirements
    3. AI analyzes and optimizes your resume
    4. Download in PDF or DOCX format
    
    **Protected:**
    - ✅ Education
    - ✅ Projects
    - ✅ Personal Info
    
    **Optimized:**
    - 🔄 Summary
    - 🔄 Skills
    - 🔄 Experience descriptions
    """)
    
    st.markdown("---")
    st.info("💡 Your education and projects remain unchanged. Only experience descriptions and skills are optimized.")

# Main content area with tabs
tab1, tab2 = st.tabs(["📤 Upload & Optimize", "📊 Preview & Download"])

with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="step-container">', unsafe_allow_html=True)
        st.subheader("📄 Step 1: Upload Your Resume")
        uploaded_file = st.file_uploader(
            "Upload your current resume", 
            type=['pdf', 'docx'],
            help="Supported formats: PDF, DOCX"
        )
        
        if uploaded_file:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            file_type = uploaded_file.name.split('.')[-1].lower()
            
            # Extract text based on file type
            if file_type == 'pdf':
                resume_text = extract_text_from_pdf(uploaded_file)
            elif file_type == 'docx':
                resume_text = extract_text_from_docx(uploaded_file)
            
            if resume_text:
                with st.expander("👁️ View extracted text (preview)"):
                    st.text(resume_text[:500] + "..." if len(resume_text) > 500 else resume_text)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="step-container">', unsafe_allow_html=True)
        st.subheader("🎯 Step 2: Job Requirements")
        job_requirements = st.text_area(
            "Paste the job description or requirements",
            height=300,
            placeholder="Paste the complete job description here, including:\n- Required skills\n- Responsibilities\n- Qualifications\n- Experience needed\n- Any specific requirements"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="step-container">', unsafe_allow_html=True)
    st.subheader("🚀 Step 3: Generate Optimized Resume")
    
    col_a, col_b, col_c = st.columns([1, 2, 1])
    
    with col_b:
        optimize_button = st.button("✨ Optimize My Resume", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    if optimize_button:
        if not api_key:
            st.error("❌ Please enter your Google API Key in the sidebar")
        elif not uploaded_file:
            st.error("❌ Please upload your resume")
        elif not job_requirements:
            st.error("❌ Please provide job requirements")
        elif not resume_text:
            st.error("❌ Could not extract text from resume. Please try another file.")
        else:
            try:
                with st.spinner("🔍 Analyzing your resume..."):
                    # Initialize LLM
                    llm = ChatGoogleGenerativeAI(
                        model="gemini-2.0-flash",
                        google_api_key=api_key,
                        temperature=0.7
                    )
                    
                    # Step 1: Extract resume data
                    extraction_chain = extraction_prompt | llm
                    
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("📋 Extracting resume information...")
                    progress_bar.progress(25)
                    
                    extraction_response = extraction_chain.invoke({"resume_text": resume_text})
                    
                    # Parse extracted data
                    extracted_text = extraction_response.content
                    if "```json" in extracted_text:
                        extracted_text = extracted_text.split("```json")[1].split("```")[0]
                    elif "```" in extracted_text:
                        extracted_text = extracted_text.split("```")[1].split("```")[0]
                    
                    extracted_data = json.loads(extracted_text.strip())
                    
                    progress_bar.progress(50)
                    status_text.text("🎯 Optimizing resume for job requirements...")
                    
                    time.sleep(1)  # Rate limiting
                    
                    # Step 2: Optimize resume
                    optimization_chain = optimization_prompt | llm
                    
                    optimization_response = optimization_chain.invoke({
                        "extracted_resume": json.dumps(extracted_data, indent=2),
                        "job_requirements": job_requirements
                    })
                    
                    # Parse optimized data
                    optimized_text = optimization_response.content
                    if "```json" in optimized_text:
                        optimized_text = optimized_text.split("```json")[1].split("```")[0]
                    elif "```" in optimized_text:
                        optimized_text = optimized_text.split("```")[1].split("```")[0]
                    
                    optimized_data = json.loads(optimized_text.strip())
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Optimization complete!")
                    
                    # Store in session state
                    st.session_state.optimized_resume = optimized_data
                    st.session_state.original_resume = extracted_data
                    
                    time.sleep(0.5)
                    st.balloons()
                    
                    st.markdown('<div class="success-box">', unsafe_allow_html=True)
                    st.success("🎉 Your resume has been successfully optimized! Switch to the 'Preview & Download' tab to view and download.")
                    st.markdown('</div>', unsafe_allow_html=True)
                    
            except json.JSONDecodeError as e:
                st.error(f"❌ Error parsing response: {str(e)}")
                with st.expander("🔍 Debug Information"):
                    st.code(extracted_text if 'extracted_text' in locals() else optimized_text)
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                if "429" in str(e):
                    st.info("⏳ Rate limit reached. Please wait a moment and try again.")

with tab2:
    if st.session_state.optimized_resume:
        resume_data = st.session_state.optimized_resume
        
        st.subheader("📋 Your Optimized Resume")
        
        # Comparison view
        if 'original_resume' in st.session_state:
            with st.expander("🔄 View Changes Made"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Original Summary:**")
                    st.info(st.session_state.original_resume.get('summary', 'N/A'))
                
                with col2:
                    st.markdown("**Optimized Summary:**")
                    st.success(resume_data.get('summary', 'N/A'))
        
        # Preview sections
        with st.container():
            st.markdown("### 👤 Personal Information")
            st.write(f"**Name:** {resume_data.get('name', 'N/A')}")
            st.write(f"**Email:** {resume_data.get('email', 'N/A')}")
            st.write(f"**Phone:** {resume_data.get('phone', 'N/A')}")
        
        with st.container():
            st.markdown("### 💼 Professional Summary")
            st.write(resume_data.get('summary', 'N/A'))
        
        with st.container():
            st.markdown("### 🛠️ Skills")
            if resume_data.get('skills'):
                skills_cols = st.columns(3)
                for idx, skill in enumerate(resume_data['skills']):
                    with skills_cols[idx % 3]:
                        st.markdown(f"✓ {skill}")
        
        with st.container():
            st.markdown("### 💼 Experience")
            for exp in resume_data.get('experience', []):
                with st.expander(f"{exp.get('title', 'Position')} at {exp.get('company', 'Company')}"):
                    st.write(f"**Duration:** {exp.get('duration', 'N/A')}")
                    st.write("**Responsibilities:**")
                    for resp in exp.get('responsibilities', []):
                        st.write(f"• {resp}")
        
        # Download section
        st.markdown("---")
        st.markdown('<div class="download-section">', unsafe_allow_html=True)
        st.subheader("📥 Download Your Resume")
        
        col1, col2 = st.columns(2)
        
        with col1:
            try:
                pdf_buffer = generate_stylish_pdf(resume_data)
                
                st.download_button(
                    label="📄 Download as PDF",
                    data=pdf_buffer,
                    file_name=f"optimized_resume_{resume_data.get('name', 'candidate').replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary"
                )
                
            except Exception as e:
                st.error(f"Error generating PDF: {str(e)}")
        
        with col2:
            try:
                docx_buffer = generate_stylish_docx(resume_data)
                
                st.download_button(
                    label="📝 Download as DOCX",
                    data=docx_buffer,
                    file_name=f"optimized_resume_{resume_data.get('name', 'candidate').replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
                
            except Exception as e:
                st.error(f"Error generating DOCX: {str(e)}")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.success("✅ Your resume is ready for download in both formats!")
    
    else:
        st.info("👆 Upload your resume and job requirements in the 'Upload & Optimize' tab to get started!")

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #7f8c8d; padding: 2rem 0;'>
        <p>🚀 Built with Streamlit, LangChain & Google Gemini AI</p>
        <p style='font-size: 0.8rem;'>Your privacy matters: We don't store your resume data</p>
    </div>
    """,
    unsafe_allow_html=True
)